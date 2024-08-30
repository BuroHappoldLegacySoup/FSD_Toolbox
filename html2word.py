from dataclasses import dataclass
from typing import List, Optional
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION
import os
import re
from PIL import Image
import io
from info import TableInfo, ImageInfo

class HTMLToWordConverter:
    """
    A class to convert HTML tables to Word document tables.
    
    This class provides functionality to extract tables from HTML content,
    convert them to Word tables, and save them in a Word document.
    """

    def __init__(self, doc_path: str, html_path: str):
        """
        Initialize the converter with an existing Word document.
        
        Args:
            doc_path (str): The path to the existing Word document.
        """
        self.doc = Document(doc_path)
        self.html_path = html_path
        self.data_folder = f"{os.path.splitext(html_path)[0]}_data"
        self.image_files: List[str] = []
        self.images: List[ImageInfo] = []
        with open(html_path, 'r', encoding='utf-8') as file:
            self.soup = BeautifulSoup(file, 'html.parser')

    def extract_image_files(self) -> None:
        self.image_files = [f for f in os.listdir(self.data_folder) if f.endswith('.png')]
        print(f"Found {len(self.image_files)} image files")

    def extract_all_tables(self):
        """
        Extract all tables from the HTML file and add them to the Word document.
        """
        tables = self.soup.find_all('table')
        for i, table in enumerate(tables):
            # Find the nearest preceding heading
            heading = table.find_previous(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])

            title = re.sub(r'^\d+(\.\d+)*\s*', '', heading.get_text(strip=True))
            
            # Create a new BeautifulSoup object with just this table
            table_html = str(table)
            self.create_word_table_from_html(table_html, title)
            
            # Add some space after each table
            self.doc.add_paragraph()

    def extract_captions(self) -> None:
        with open(self.html_path, 'r', encoding='utf-8') as file:
            soup = BeautifulSoup(file, 'html.parser')
            img_tags = soup.find_all('img')
            for img in img_tags[2:]:
                src = img.get('src')
                if src:
                    file_name = os.path.basename(src)
                    h2 = img.find_previous('h2')
                    if h2:
                        caption = re.sub(r'^[\d.]+ ', '', h2.text.strip())
                        self.images.append(ImageInfo(file_name, caption))
        print(f"Extracted {len(self.images)} image captions")

    def add_images_to_word_document(self) -> None:
        """
        Add images and captions to an existing Word document.
        """   
        for image in self.images:
            if image.filename in self.image_files:
                # Add page break before new content
                self.doc.add_page_break()
                new_section = self.doc.add_section()

                # Set properties for the new section
                new_section.orientation = WD_ORIENTATION.LANDSCAPE
                new_section.page_width = Inches(11)  # Standard letter size
                new_section.page_height = Inches(8.5)
                new_section.left_margin = Inches(1)
                new_section.right_margin = Inches(1)
                
                img_path = os.path.join(self.data_folder, image.filename)
                """with Image.open(img_path) as img:
                    # Rotate the image 90 degrees counterclockwise
                    rotated_img = img.rotate(90, expand=True)
                    
                    # Save the rotated image to a bytes buffer
                    img_buffer = io.BytesIO()
                    rotated_img.save(img_buffer, format=img.format)
                    img_buffer.seek(0)"""
                
                # Add rotated image
                
                self.doc.add_picture(img_path, width=Inches(9))
                last_paragraph = self.doc.paragraphs[-1] 
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Add caption
                p = self.doc.add_paragraph()
                p.alignment = 1  # Center alignment
                p.add_run(image.caption).italic = True

        # Add final portrait section
        final_section = self.doc.add_section()
        final_section.orientation = WD_ORIENTATION.PORTRAIT
        final_section.page_width = Inches(8.5)  # Standard letter size
        final_section.page_height = Inches(11)
        final_section.left_margin = Inches(1)
        final_section.right_margin = Inches(1)

    @staticmethod
    def extract_table_after_heading(soup: BeautifulSoup, main_title: str, heading_text: str) -> str:
        # Find the exact match for the main title (h1)
        main_heading = soup.find('h1', string=lambda text: text and main_title in text)
        if main_heading:
            # Find the next h2 that contains the heading_text
            heading = main_heading.find_next('h2', string=lambda text: text and heading_text in text)
            if heading:
                # Check if the next h2 is a subheading
                next_h2 = heading.find_next('h2')
                if next_h2:
                    # Extract all content between the current h2 and the next h2
                    content = []
                    for sibling in heading.next_siblings:
                        if sibling == next_h2:
                            break
                        content.append(str(sibling))
                    content_str = ''.join(content)
                    # Find the first table in this content
                    soup_content = BeautifulSoup(content_str, 'html.parser')
                    target_table = soup_content.find('table')
                    if target_table:
                        return str(target_table)
        return None

    @staticmethod
    def apply_cell_formatting(cell, color_hex: str):
        """
        Apply background color to a cell.
        
        Args:
            cell: The Word table cell to format.
            color_hex (str): The hex color code to apply as background.
        """
        if color_hex:
            cell_xml = cell._element
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), color_hex)
            cell_xml.get_or_add_tcPr().append(shading)

    def create_word_table_from_html(self, html_content: str, title: str):
        soup = BeautifulSoup(html_content, 'html.parser')
        table = soup.find('table')
        if not table:
            print(f"No table found for title: {title}")
            return
        self.doc.add_heading(title, level=1)
        headers = table.find_all('tr')[0].find_all(['th', 'td'])
        max_columns = sum(int(th.get('colspan', 1)) for th in headers)
        rows = table.find_all('tr')
        num_rows = len(rows)
        word_table = self.doc.add_table(rows=num_rows, cols=max_columns)
        word_table.style = 'Table Grid'
        self._fill_table_content(word_table, rows, max_columns)
        self._remove_empty_columns(word_table)
        self._remove_empty_rows(word_table)
        self._apply_row_colors(word_table)
        #print(f"Created table: {title}")

    def _fill_table_content(self, word_table, rows, max_columns):
        """
        Fill the Word table with content from the HTML table.
        
        Args:
            word_table: The Word table to fill.
            rows: The rows from the HTML table.
            max_columns (int): The maximum number of columns in the table.
        """
        for row_idx, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            col_idx = 0
            while cells and col_idx < max_columns:
                cell = cells.pop(0)
                colspan = int(cell.get('colspan', 1))
                cell_text = cell.get_text(strip=True)
                cell_style = cell.get('style', '')
                cell_color = None
                if 'background-color:' in cell_style:
                    color_hex = cell_style.split('background-color:')[1].split(';')[0].strip()
                    if color_hex.startswith('#'):
                        cell_color = color_hex[1:]
                self._fill_cell_content(word_table, row_idx, col_idx, cell_text, cell_color)
                col_idx += colspan

    def _fill_cell_content(self, word_table, row: int, col: int, content: str, color_hex: str):
        """
        Fill the content and background color of a cell in the Word table.
        
        Args:
            word_table: The Word table containing the cell.
            row (int): The row index of the cell.
            col (int): The column index of the cell.
            content (str): The text content to fill in the cell.
            color_hex (str): The hex color code for the cell background.
        """
        word_cell = word_table.cell(row, col)
        word_cell.text = content
        if color_hex:
            self.apply_cell_formatting(word_cell, color_hex)

    def _remove_empty_columns(self, word_table):
        """
        Remove empty columns from the Word table.
        
        Args:
            word_table: The Word table to process.
        """
        columns_to_remove = [i for i in range(len(word_table.columns)) if self._is_column_empty(word_table, i)]
        for col_idx in reversed(columns_to_remove):
            for row in word_table.rows:
                cell = row.cells[col_idx]
                cell._element.getparent().remove(cell._element)

    @staticmethod
    def _is_column_empty(word_table, col_idx: int) -> bool:
        """
        Check if a column is entirely empty.
        
        Args:
            word_table: The Word table to check.
            col_idx (int): The index of the column to check.
        
        Returns:
            bool: True if the column is empty, False otherwise.
        """
        return all(row.cells[col_idx].text.strip() == '' for row in word_table.rows)

    def _remove_empty_rows(self, word_table):
        """
        Remove empty rows from the Word table.
        
        Args:
            word_table: The Word table to process.
        """
        rows_to_remove = [row for row in word_table.rows if self._is_row_empty(row)]
        for row in rows_to_remove:
            tbl = word_table._tbl
            tbl.remove(row._tr)

    @staticmethod
    def _is_row_empty(row) -> bool:
        """
        Check if a row is entirely empty.
        
        Args:
            row: The row to check.
        
        Returns:
            bool: True if the row is empty, False otherwise.
        """
        return all(cell.text.strip() == '' for cell in row.cells)

    def _delete_last_page_in_template(self):
        for element in reversed(self.doc.element.body):
            if element.tag.endswith('sectPr'):
                self.doc.element.body.remove(element)
                break

    def _apply_row_colors(self, word_table):
        """
        Apply the background color of the second cell to the remaining cells of each row.
        
        Args:
            word_table: The Word table to process.
        """
        for row in word_table.rows:
            if len(row.cells) > 1:
                second_cell_color = row.cells[1]._element.xpath('.//w:shd/@w:fill')
                if second_cell_color:
                    second_cell_color = second_cell_color[0]
                    for cell in row.cells[2:]:
                        self.apply_cell_formatting(cell, second_cell_color)

    def save(self, filename: str):
        """
        Save the Word document to a file.
        
        Args:
            filename (str): The name of the file to save the document to.
        """
        self.doc.save(filename)

    def process_html_file(self, table_info_list: Optional[List[TableInfo]] = None):
        """
        Process the HTML file to extract tables and add them to the Word document.
        If table_info_list is provided, extract specific tables, otherwise extract all tables.
        
        Args:
            table_info_list (Optional[List[TableInfo]]): A list of TableInfo objects specifying the tables to extract.
                                                         If None, all tables will be extracted.
        """
        if table_info_list:
            for table_info in table_info_list:
                table_html = self.extract_table_after_heading(self.soup, table_info.main_title, table_info.heading_text)
                if table_html:
                    self.create_word_table_from_html(table_html, table_info.title)
                else:
                    print(f"No table found for heading: {table_info.heading_text}")
        else:
            self.extract_all_tables()

    

# Example usage
if __name__ == "__main__":
    converter = HTMLToWordConverter(r"C:\Users\vmylavarapu\Desktop\Template.docx", r"C:\Users\vmylavarapu\Desktop\FSRG\pr1.html")

    # Define the tables to extract
    """table_info_list_1 = [
        TableInfo('1.1 Materials', 'Materials'),
        TableInfo('1.6 Surfaces', 'Surfaces'),
        TableInfo('2.1 Nodal Supports', 'Nodal Supports')
    ]

    table_info_list_2 = [
        TableInfo('1.2 Sections', 'Cross sections'),
        TableInfo('4.1 Load Cases', 'Load Cases')
    ]"""

    # Process the HTML file
    table_info_list = [
        TableInfo('Basic Objects', 'Materials', 'Materials'),
        TableInfo('Basic Objects', 'Surfaces', 'Surfaces'),
        TableInfo('Types for Nodes', 'Nodal Supports', 'Nodal Supports'),
        TableInfo('Types for Lines', 'Line Supports', 'Line Supports'),
        TableInfo('Basic Objects', 'Sections', 'Sections'),
        TableInfo('Basic Objects', 'Materials', 'Members'),
        ]
    converter._delete_last_page_in_template()
    #converter.process_html_file(table_info_list_1)
    converter.process_html_file()
    converter.extract_image_files()
    converter.extract_captions()
    converter.add_images_to_word_document()
    #converter.process_html_file(table_info_list_2)

    # Save the Word document
    converter.save(r"C:\Users\vmylavarapu\Downloads\Output.docx")
    print("The Word document 'Output.docx' was successfully created.")