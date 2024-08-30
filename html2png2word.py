import os
import re
from dataclasses import dataclass
from typing import List
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup

@dataclass
class ImageInfo:
    """Stores information about an image."""
    filename: str
    caption: str

class HTMLImageExtractor:
    """
    A class to extract images and captions from an HTML file and add them to an existing Word document.

    This class ignores the first two PNG files mentioned in the HTML file and 
    extracts captions for the remaining images from the preceding <h2> tags.

    Attributes:
        html_file (str): Path to the input HTML file.
        data_folder (str): Path to the folder containing the image files.
        image_files (List[str]): List of PNG filenames in the data folder.
        images (List[ImageInfo]): List of ImageInfo objects containing filename and caption.
    """

    def __init__(self, html_file: str):
        """
        Initialize the HTMLImageExtractor.

        Args:
            html_file (str): Path to the input HTML file.
        """
        self.html_file = html_file
        self.data_folder = f"{os.path.splitext(html_file)[0]}_data"
        self.image_files: List[str] = []
        self.images: List[ImageInfo] = []

    def extract_image_files(self) -> None:
        """Extract all PNG filenames from the data folder."""
        self.image_files = [f for f in os.listdir(self.data_folder) if f.endswith('.png')]

    def extract_captions(self) -> None:
        """
        Extract image captions from the HTML file.

        This method ignores the first two PNG files mentioned in the HTML and
        extracts captions for the remaining images from the preceding <h2> tags.
        """
        with open(self.html_file, 'r', encoding='utf-8') as file:
            soup = BeautifulSoup(file, 'html.parser')
            img_tags = soup.find_all('img')
            
            # Ignore the first two PNG files
            for img in img_tags[2:]:
                src = img.get('src')
                if src:
                    file_name = os.path.basename(src)
                    h2 = img.find_previous('h2')
                    if h2:
                        caption = re.sub(r'^[\d.]+ ', '', h2.text.strip())
                        self.images.append(ImageInfo(file_name, caption))

    def add_images_to_word_document(self, word_file: str) -> None:
        """
        Add images and captions to an existing Word document.

        Args:
            word_file (str): Path to the existing Word document.
        """
        doc = Document(word_file)
        
        for image in self.images:
            if image.filename in self.image_files:
                # Add page break before new content
                doc.add_page_break()
                
                # Add image
                doc.add_picture(os.path.join(self.data_folder, image.filename), width=Inches(6))
                
                # Add caption
                p = doc.add_paragraph()
                p.alignment = 1  # Center alignment
                p.add_run(image.caption).italic = True
        
        #doc.save(word_file)

    def process(self, word_file: str) -> None:
        """
        Process the HTML file and add images to the existing Word document.

        This method calls the other methods in the correct order to complete the task.

        Args:
            word_file (str): Path to the existing Word document.
        """
        self.extract_image_files()
        self.extract_captions()
        self.add_images_to_word_document(word_file)

# Usage example
if __name__ == "__main__":
    extractor = HTMLImageExtractor(r"C:\Users\vmylavarapu\Desktop\test template.html")
    extractor.process(r"C:\Users\vmylavarapu\Desktop\FSD\output_report.docx")