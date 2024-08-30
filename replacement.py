import os
from docx import Document
import zipfile
import re
import tempfile
from typing import List, Dict
from dataclasses import dataclass, field

@dataclass
class WordReplacement:
    """
    Represents a word replacement pair.

    Attributes:
        old_word (str): The word to be replaced.
        new_word (str): The replacement word.
    """
    old_word: str
    new_word: str

@dataclass
class DocumentWordReplacer:
    """
    A class for replacing multiple words in a Microsoft Word document.

    This class provides functionality to replace specified words throughout a Word document,
    including in paragraphs, tables, headers, footers, and XML content.

    Attributes:
        file_path (str): The path to the Word document to be modified.
        replacements (List[WordReplacement]): A list of WordReplacement objects.
    """
    file_path: str
    replacements: List[WordReplacement] = field(default_factory=list)

    def add_replacement(self, old_word: str, new_word: str):
        """
        Add a new word replacement pair.

        Args:
            old_word (str): The word to be replaced.
            new_word (str): The replacement word.
        """
        self.replacements.append(WordReplacement(old_word, new_word))

    def _replace_in_docx(self, doc: Document):
        """
        Replace words in the Document object using python-docx.

        This method replaces words in paragraphs, tables, headers, and footers.

        Args:
            doc (Document): A python-docx Document object.
        """
        for paragraph in doc.paragraphs:
            for replacement in self.replacements:
                if replacement.old_word in paragraph.text:
                    for run in paragraph.runs:
                        if replacement.old_word in run.text:
                            run.text = run.text.replace(replacement.old_word, replacement.new_word)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for replacement in self.replacements:
                        if replacement.old_word in cell.text:
                            cell.text = cell.text.replace(replacement.old_word, replacement.new_word)

        for section in doc.sections:
            for header in section.header.paragraphs:
                for replacement in self.replacements:
                    if replacement.old_word in header.text:
                        for run in header.runs:
                            if replacement.old_word in run.text:
                                run.text = run.text.replace(replacement.old_word, replacement.new_word)
            
            for footer in section.footer.paragraphs:
                for replacement in self.replacements:
                    if replacement.old_word in footer.text:
                        for run in footer.runs:
                            if replacement.old_word in run.text:
                                run.text = run.text.replace(replacement.old_word, replacement.new_word)

    def _replace_in_xml(self, temp_path: str, modified_path: str):
        """
        Replace words in the XML content of the document.

        This method accesses and modifies the raw XML content of the document.

        Args:
            temp_path (str): Path to the temporary file.
            modified_path (str): Path where the modified document will be saved.
        """
        with zipfile.ZipFile(temp_path, 'r') as zin:
            with zipfile.ZipFile(modified_path, 'w') as zout:
                for item in zin.infolist():
                    buffer = zin.read(item.filename)
                    if item.filename.endswith('.xml'):
                        content = buffer.decode('utf-8')
                        for replacement in self.replacements:
                            content = re.sub(re.escape(replacement.old_word), replacement.new_word, content)
                        buffer = content.encode('utf-8')
                    zout.writestr(item, buffer)

    def replace_words(self, folder_path) -> str:
        """
        Perform the word replacement process on the document.

        This method orchestrates the entire replacement process, including
        creating temporary files, calling other methods to perform replacements,
        and cleaning up temporary files.

        Args:
            folder_path (str): The path to the folder where the modified document will be saved.

        Returns:
            str: The path to the modified document.

        Raises:
            FileNotFoundError: If the specified Word document does not exist.
        """
        if not os.path.exists(self.file_path):
            raise FileNotFoundError(f"The file {self.file_path} does not exist.")

        doc = Document(self.file_path)
        self._replace_in_docx(doc)

        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            temp_path = tmp_file.name

        doc.save(temp_path)

        filename = os.path.basename(self.file_path)
        modified_filename = 'Report_output.docx'
        modified_path = os.path.join(folder_path, modified_filename)

        self._replace_in_xml(temp_path, modified_path)

        os.remove(temp_path)

        return modified_path

# Usage example
if __name__ == "__main__":
    file_path = r'C:\Users\vmylavarapu\Desktop\Template.docx'
    replacer = DocumentWordReplacer(file_path)
    
    replacer.add_replacement('Projekttitel', 'new')
    replacer.add_replacement('Berichttitel', 'replacement1')
    replacer.add_replacement('XXXX-BHE-XX-XX-XX-X-XXXX', 'replacement2')
    
    modified_file_path = replacer.replace_words()